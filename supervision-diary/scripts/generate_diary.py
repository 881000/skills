#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
监理日志生成器
基于模板生成监理日志docx文件，具备内容优化表达、智能摘要精简、
段落编号、行距控制、首行缩进等专业排版功能。
"""

import argparse
import os
import sys
import json
import datetime
import re
import urllib.request
import urllib.parse
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


WEEKDAY_NAMES = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']

# 当日监理工作情况未提供时的通用表述（不允许留空）
DEFAULT_SUPERVISION = ('对施工现场进行巡视检查，检查施工质量、安全生产及文明施工情况，'
                       '各工序施工质量符合设计及规范要求。对关键部位、关键工序实施旁站监理，'
                       '相关施工资料同步完善。')

# 四大板块排版限制配置
SECTION_LIMITS = {
    'construction': {'max_chars': 600, 'max_lines': 18, 'label': '当日施工进展情况'},
    'supervision':  {'max_chars': 260, 'max_lines': 6,  'label': '当日监理工作情况'},
    'problem':      {'max_chars': 160, 'max_lines': 4,  'label': '当日存在的问题及处理情况'},
    'other':        {'max_chars': 120, 'max_lines': 3,  'label': '其他有关事项'},
}

# 排版参数（用于物理行数估算）
CELL_WIDTH_TWIPS = 9888       # 内容单元格总宽度（15列合计）
CELL_MARGIN_TWIPS = 240       # 单元格左右内边距合计（保守估算，确保不超页）
FONT_SIZE_PT = 10.5            # 五号字
CHAR_WIDTH_PT = FONT_SIZE_PT   # 中文字符宽度约等于字号
FIRST_LINE_INDENT_CHARS = 2    # 首行缩进2字符
NUMBER_PREFIX_CHARS = 2         # 编号"1、"占2字符

# 各板块关键词（用于摘要时判断句子重要性）
SECTION_KEYWORDS = {
    'construction': [
        '浇筑', '绑扎', '安装', '拆除', '开挖', '回填', '完成', '进行', '正在',
        '楼', '层', '段', '区', '轴', '方量', 'm³', '立方', '工程量', '进度',
        '钢筋', '模板', '混凝土', '砌体', '抹灰', '防水', '保温', '装饰', '装修',
        '水电', '预埋', '管线', '脚手架', '支撑', '养护', '放线', '测量',
    ],
    'supervision': [
        '巡视', '检查', '旁站', '验收', '审核', '审查', '签署', '签发',
        '会议', '例会', '协调', '要求', '整改', '合格', '符合', '不符合',
        '质量', '安全', '进度', '投资', '合同', '资料', '报审', '报验',
        '坍落度', '试块', '取样', '检测', '复核', '见证',
    ],
    'problem': [
        '问题', '隐患', '不符合', '不合格', '整改', '责令', '要求', '复查',
        '合格', '验收', '通知', '停工', '返工', '纠正', '预防', '处理',
        '安全', '质量', '规范', '方案', '设计', '标准',
    ],
    'other': [
        '建设单位', '设计单位', '勘察单位', '施工单位', '监理单位',
        '检查', '指导', '调研', '会议', '通知', '文件', '来访', '接待',
    ],
}


def parse_date(date_str):
    """解析日期字符串，支持多种格式，返回 datetime.date 对象。"""
    if not date_str:
        return None
    date_str = date_str.strip()
    formats = ['%Y年%m月%d日', '%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d', '%Y%m%d']
    for fmt in formats:
        try:
            return datetime.datetime.strptime(date_str, fmt).date()
        except ValueError:
            continue
    match = re.search(r'(\d{4})\D*(\d{1,2})\D*(\d{1,2})', date_str)
    if match:
        try:
            year, month, day = map(int, match.groups())
            return datetime.date(year, month, day)
        except (ValueError, TypeError):
            pass
    return None


def get_weekday(date_str):
    """根据日期字符串返回中文星期几。"""
    dt = parse_date(date_str)
    if dt is None:
        return ''
    return WEEKDAY_NAMES[dt.weekday()]


# ============ 天气/气温自动获取（基于城市） ============

# WMO 天气代码 -> 中文天气描述（Open-Meteo daily.weather_code）
WMO_WEATHER = {
    0: '晴', 1: '晴', 2: '多云', 3: '阴',
    45: '雾', 48: '雾凇',
    51: '毛毛雨', 53: '毛毛雨', 55: '毛毛雨',
    56: '冻雨', 57: '冻雨',
    61: '小雨', 63: '中雨', 65: '大雨',
    66: '冻雨', 67: '冻雨',
    71: '小雪', 73: '中雪', 75: '大雪', 77: '雪粒',
    80: '阵雨', 81: '阵雨', 82: '强阵雨',
    85: '阵雪', 86: '阵雪',
    95: '雷阵雨', 96: '雷阵雨伴冰雹', 99: '雷阵雨伴冰雹',
}



CITY_COORDS = {
    '七台河': {'lat': 45.768, 'lon': 130.9953},
    '三亚': {'lat': 18.25435, 'lon': 109.50947},
    '三明': {'lat': 26.24861, 'lon': 117.61861},
    '三门峡': {'lat': 34.78081, 'lon': 111.19287},
    '上海': {'lat': 31.22222, 'lon': 121.45806},
    '上饶': {'lat': 28.45179, 'lon': 117.94287},
    '东莞': {'lat': 23.01797, 'lon': 113.74866},
    '东营': {'lat': 37.46271, 'lon': 118.49165},
    '中卫': {'lat': 37.51129, 'lon': 105.19067},
    '中山': {'lat': 22.52306, 'lon': 113.37912},
    '临汾': {'lat': 36.08889, 'lon': 111.51889},
    '临沂': {'lat': 35.06306, 'lon': 118.34278},
    '临沧': {'lat': 23.87972, 'lon': 100.09455},
    '丹东': {'lat': 40.12917, 'lon': 124.39472},
    '丽水': {'lat': 28.46042, 'lon': 119.91029},
    '丽江': {'lat': 26.86879, 'lon': 100.22072},
    '乌兰察布': {'lat': 40.993, 'lon': 113.133},
    '乌海': {'lat': 39.68442, 'lon': 106.81583},
    '乌鲁木齐': {'lat': 43.80096, 'lon': 87.60046},
    '乐山': {'lat': 29.56227, 'lon': 103.76386},
    '九江': {'lat': 29.70475, 'lon': 116.00206},
    '云浮': {'lat': 22.92787, 'lon': 112.03809},
    '亳州': {'lat': 33.87722, 'lon': 115.77028},
    '伊宁': {'lat': 43.91515, 'lon': 81.32151},
    '伊春': {'lat': 47.72143, 'lon': 128.87529},
    '佛山': {'lat': 23.02677, 'lon': 113.13148},
    '佳木斯': {'lat': 46.79711, 'lon': 130.31117},
    '保定': {'lat': 38.87288, 'lon': 115.46246},
    '保山': {'lat': 25.11626, 'lon': 99.16366},
    '信阳': {'lat': 32.12278, 'lon': 114.06556},
    '儋州': {'lat': 19.52134, 'lon': 109.57895},
    '克拉玛依': {'lat': 45.58473, 'lon': 84.88724},
    '六安': {'lat': 31.62796, 'lon': 116.27582},
    '六盘水': {'lat': 26.59444, 'lon': 104.83333},
    '兰州': {'lat': 36.05701, 'lon': 103.83987},
    '内江': {'lat': 29.58354, 'lon': 105.06216},
    '包头': {'lat': 40.6516, 'lon': 109.84389},
    '北京': {'lat': 39.9075, 'lon': 116.39723},
    '北海': {'lat': 21.48349, 'lon': 109.11549},
    '十堰': {'lat': 32.6475, 'lon': 110.77806},
    '南京': {'lat': 32.06167, 'lon': 118.77778},
    '南充': {'lat': 30.79508, 'lon': 106.08473},
    '南宁': {'lat': 22.81667, 'lon': 108.31667},
    '南平': {'lat': 26.645, 'lon': 118.17361},
    '南昌': {'lat': 28.68396, 'lon': 115.85306},
    '南通': {'lat': 32.03028, 'lon': 120.87472},
    '南阳': {'lat': 33.00524, 'lon': 112.54659},
    '厦门': {'lat': 24.47979, 'lon': 118.08187},
    '双鸭山': {'lat': 46.67686, 'lon': 131.13274},
    '台中': {'lat': 24.1469, 'lon': 120.6839},
    '台北': {'lat': 25.05306, 'lon': 121.52639},
    '台南': {'lat': 22.99083, 'lon': 120.21333},
    '台州': {'lat': 28.66266, 'lon': 121.43312},
    '合肥': {'lat': 31.86389, 'lon': 117.28083},
    '吉安': {'lat': 22.58333, 'lon': 113.08333},
    '吉林': {'lat': 43.84652, 'lon': 126.5608},
    '吐鲁番': {'lat': 42.94769, 'lon': 89.17886},
    '吕梁': {'lat': 37.5192, 'lon': 111.14436},
    '吴忠': {'lat': 37.9867, 'lon': 106.201},
    '周口': {'lat': 33.63333, 'lon': 114.63333},
    '呼伦贝尔': {'lat': 49.21141, 'lon': 119.75582},
    '呼和浩特': {'lat': 40.81056, 'lon': 111.65222},
    '和田': {'lat': 37.10927, 'lon': 79.93433},
    '咸宁': {'lat': 29.84347, 'lon': 114.32201},
    '咸阳': {'lat': 34.33778, 'lon': 108.70261},
    '哈密': {'lat': 42.83393, 'lon': 93.50601},
    '哈尔滨': {'lat': 45.75, 'lon': 126.65},
    '唐山': {'lat': 39.64381, 'lon': 118.18319},
    '商丘': {'lat': 34.4143, 'lon': 115.65613},
    '商洛': {'lat': 33.86667, 'lon': 109.93056},
    '喀什': {'lat': 39.46718, 'lon': 75.98675},
    '嘉兴': {'lat': 30.7522, 'lon': 120.75},
    '嘉峪关': {'lat': 39.81121, 'lon': 98.28618},
    '四平': {'lat': 43.16143, 'lon': 124.37785},
    '固原': {'lat': 36.00667, 'lon': 106.28083},
    '塔城': {'lat': 46.74535, 'lon': 82.95847},
    '大同': {'lat': 40.09361, 'lon': 113.29139},
    '大庆': {'lat': 46.58333, 'lon': 125.0},
    '大连': {'lat': 38.91222, 'lon': 121.60222},
    '天水': {'lat': 34.57952, 'lon': 105.74238},
    '天津': {'lat': 39.14222, 'lon': 117.17667},
    '太原': {'lat': 37.86944, 'lon': 112.56028},
    '威海': {'lat': 37.50914, 'lon': 122.11356},
    '娄底': {'lat': 27.73444, 'lon': 111.99444},
    '孝感': {'lat': 30.92689, 'lon': 113.92221},
    '宁德': {'lat': 26.66167, 'lon': 119.52278},
    '宁波': {'lat': 29.87819, 'lon': 121.54945},
    '安庆': {'lat': 30.51365, 'lon': 117.04723},
    '安康': {'lat': 32.68, 'lon': 109.01722},
    '安阳': {'lat': 36.096, 'lon': 114.38278},
    '安顺': {'lat': 26.25, 'lon': 105.93333},
    '定西': {'lat': 35.57088, 'lon': 104.62303},
    '宜宾': {'lat': 28.7593, 'lon': 104.63994},
    '宜昌': {'lat': 30.71444, 'lon': 111.28472},
    '宜春': {'lat': 27.83333, 'lon': 114.4},
    '宝鸡': {'lat': 34.36775, 'lon': 107.23705},
    '宣城': {'lat': 30.9525, 'lon': 118.75528},
    '宿州': {'lat': 33.63611, 'lon': 116.97889},
    '宿迁': {'lat': 33.94917, 'lon': 118.29583},
    '山南': {'lat': 28.42495, 'lon': 120.92872},
    '岳阳': {'lat': 29.37455, 'lon': 113.09481},
    '崇左': {'lat': 22.38161, 'lon': 107.3683},
    '巴中': {'lat': 31.8694, 'lon': 106.74432},
    '巴彦淖尔': {'lat': 40.9264, 'lon': 107.739},
    '常州': {'lat': 31.77359, 'lon': 119.95401},
    '常德': {'lat': 29.03205, 'lon': 111.69844},
    '平凉': {'lat': 35.53917, 'lon': 106.68611},
    '平顶山': {'lat': 33.73091, 'lon': 113.31554},
    '广元': {'lat': 32.44201, 'lon': 105.823},
    '广安': {'lat': 29.91159, 'lon': 116.92857},
    '广州': {'lat': 23.11667, 'lon': 113.25},
    '庆阳': {'lat': 35.70976, 'lon': 107.64455},
    '库尔勒': {'lat': 41.76055, 'lon': 86.15231},
    '廊坊': {'lat': 39.52079, 'lon': 116.71471},
    '延吉': {'lat': 42.88825, 'lon': 129.50241},
    '延安': {'lat': 27.055, 'lon': 109.026},
    '开封': {'lat': 34.7986, 'lon': 114.30742},
    '张家口': {'lat': 40.78341, 'lon': 114.87139},
    '张家界': {'lat': 29.12944, 'lon': 110.47833},
    '张掖': {'lat': 38.93417, 'lon': 100.45167},
    '徐州': {'lat': 34.20442, 'lon': 117.28386},
    '德州': {'lat': 37.44661, 'lon': 116.36706},
    '德阳': {'lat': 31.13019, 'lon': 104.38198},
    '忻州': {'lat': 38.40917, 'lon': 112.73333},
    '怀化': {'lat': 27.56337, 'lon': 110.00404},
    '惠州': {'lat': 23.11147, 'lon': 114.41523},
    '成都': {'lat': 30.66667, 'lon': 104.06667},
    '扬州': {'lat': 32.39722, 'lon': 119.43583},
    '承德': {'lat': 40.9519, 'lon': 117.95883},
    '抚州': {'lat': 27.95999, 'lon': 116.33333},
    '抚顺': {'lat': 41.88669, 'lon': 123.94363},
    '拉萨': {'lat': 29.65, 'lon': 91.1},
    '揭阳': {'lat': 23.5418, 'lon': 116.36581},
    '攀枝花': {'lat': 26.58509, 'lon': 101.71276},
    '新乡': {'lat': 35.19033, 'lon': 113.80151},
    '新余': {'lat': 27.80429, 'lon': 114.93335},
    '无锡': {'lat': 31.56887, 'lon': 120.28857},
    '日喀则': {'lat': 29.3519, 'lon': 89.31},
    '日照': {'lat': 35.41414, 'lon': 119.52908},
    '昆明': {'lat': 25.03889, 'lon': 102.71833},
    '昌吉': {'lat': 44.00782, 'lon': 87.30461},
    '昌都': {'lat': 24.1978, 'lon': 107.261},
    '昭通': {'lat': 27.31667, 'lon': 103.71667},
    '晋中': {'lat': 37.68403, 'lon': 112.75471},
    '晋城': {'lat': 35.50222, 'lon': 112.83278},
    '普洱': {'lat': 23.04053, 'lon': 101.03683},
    '景德镇': {'lat': 29.2947, 'lon': 117.20789},
    '曲靖': {'lat': 25.48333, 'lon': 103.78333},
    '朔州': {'lat': 39.31583, 'lon': 112.4225},
    '朝阳': {'lat': 40.48444, 'lon': 123.53694},
    '本溪': {'lat': 41.28861, 'lon': 123.765},
    '来宾': {'lat': 23.74743, 'lon': 109.22222},
    '杭州': {'lat': 30.29365, 'lon': 120.16142},
    '松原': {'lat': 45.12902, 'lon': 124.82769},
    '林芝': {'lat': 29.3033, 'lon': 94.3353},
    '枣庄': {'lat': 34.86472, 'lon': 117.55417},
    '柳州': {'lat': 24.32405, 'lon': 109.40698},
    '株洲': {'lat': 27.83333, 'lon': 113.15},
    '桂林': {'lat': 25.28022, 'lon': 110.29639},
    '梅州': {'lat': 24.28859, 'lon': 116.11768},
    '梧州': {'lat': 23.48054, 'lon': 111.28848},
    '榆林': {'lat': 38.29181, 'lon': 109.73753},
    '武威': {'lat': 37.92672, 'lon': 102.63202},
    '武汉': {'lat': 30.58333, 'lon': 114.26667},
    '毕节': {'lat': 27.30193, 'lon': 105.28627},
    '永州': {'lat': 26.42389, 'lon': 111.61306},
    '汉中': {'lat': 33.07507, 'lon': 107.02214},
    '汕头': {'lat': 23.35489, 'lon': 116.67876},
    '汕尾': {'lat': 22.78199, 'lon': 115.3475},
    '江门': {'lat': 22.58333, 'lon': 113.08333},
    '池州': {'lat': 30.66134, 'lon': 117.47783},
    '沈阳': {'lat': 41.79222, 'lon': 123.43278},
    '沧州': {'lat': 38.31124, 'lon': 116.85334},
    '河池': {'lat': 24.69285, 'lon': 108.08376},
    '河源': {'lat': 23.73333, 'lon': 114.68333},
    '泉州': {'lat': 24.91389, 'lon': 118.58583},
    '泰安': {'lat': 41.32813, 'lon': 122.45553},
    '泰州': {'lat': 32.49069, 'lon': 119.90812},
    '泸州': {'lat': 28.8903, 'lon': 105.42575},
    '洛阳': {'lat': 34.67345, 'lon': 112.43684},
    '济南': {'lat': 36.66833, 'lon': 116.99722},
    '济宁': {'lat': 35.405, 'lon': 116.58139},
    '海东': {'lat': 36.48, 'lon': 102.41639},
    '海口': {'lat': 20.03421, 'lon': 110.34651},
    '淄博': {'lat': 36.79056, 'lon': 118.06333},
    '淮北': {'lat': 33.97444, 'lon': 116.79167},
    '淮南': {'lat': 32.62639, 'lon': 116.99694},
    '淮安': {'lat': 33.58861, 'lon': 119.01917},
    '深圳': {'lat': 22.54554, 'lon': 114.0683},
    '清远': {'lat': 23.7, 'lon': 113.03333},
    '温州': {'lat': 27.99942, 'lon': 120.66682},
    '渭南': {'lat': 34.50355, 'lon': 109.50891},
    '湖州': {'lat': 30.8703, 'lon': 120.0933},
    '湘潭': {'lat': 27.85, 'lon': 112.9},
    '湛江': {'lat': 21.23391, 'lon': 110.38749},
    '滁州': {'lat': 32.32194, 'lon': 118.29778},
    '滨州': {'lat': 37.36667, 'lon': 118.01667},
    '漯河': {'lat': 33.56394, 'lon': 114.04272},
    '漳州': {'lat': 24.51333, 'lon': 117.65556},
    '潍坊': {'lat': 36.71, 'lon': 119.10194},
    '潮州': {'lat': 23.65396, 'lon': 116.62262},
    '澳门': {'lat': 22.20056, 'lon': 113.54611},
    '濮阳': {'lat': 35.75641, 'lon': 115.04363},
    '烟台': {'lat': 37.47649, 'lon': 121.44081},
    '焦作': {'lat': 35.23925, 'lon': 113.23914},
    '牡丹江': {'lat': 44.54804, 'lon': 129.62595},
    '玉林': {'lat': 22.6305, 'lon': 110.14686},
    '玉溪': {'lat': 24.355, 'lon': 102.54222},
    '珠海': {'lat': 22.27694, 'lon': 113.56778},
    '白城': {'lat': 45.61751, 'lon': 122.83302},
    '白山': {'lat': 41.93853, 'lon': 126.41965},
    '白银': {'lat': 36.54696, 'lon': 104.17023},
    '百色': {'lat': 23.89013, 'lon': 106.62684},
    '益阳': {'lat': 26.38893, 'lon': 112.37925},
    '盐城': {'lat': 33.3575, 'lon': 120.1573},
    '盘锦': {'lat': 41.121, 'lon': 122.0739},
    '眉山': {'lat': 30.04392, 'lon': 103.83696},
    '石嘴山': {'lat': 38.98082, 'lon': 106.3892},
    '石家庄': {'lat': 38.04139, 'lon': 114.47861},
    '福州': {'lat': 26.06139, 'lon': 119.30611},
    '秦皇岛': {'lat': 39.94104, 'lon': 119.58936},
    '绍兴': {'lat': 30.00237, 'lon': 120.57864},
    '绥化': {'lat': 46.64814, 'lon': 126.96656},
    '绵阳': {'lat': 31.46784, 'lon': 104.68168},
    '聊城': {'lat': 36.45064, 'lon': 116.00247},
    '肇庆': {'lat': 23.04893, 'lon': 112.46091},
    '自贡': {'lat': 29.34162, 'lon': 104.77689},
    '舟山': {'lat': 29.98869, 'lon': 122.20488},
    '芜湖': {'lat': 31.35259, 'lon': 118.42947},
    '苏州': {'lat': 31.30408, 'lon': 120.59538},
    '茂名': {'lat': 21.66625, 'lon': 110.91364},
    '荆州': {'lat': 30.35028, 'lon': 112.19028},
    '荆门': {'lat': 31.03361, 'lon': 112.20472},
    '莆田': {'lat': 25.43944, 'lon': 119.01028},
    '菏泽': {'lat': 35.23929, 'lon': 115.47358},
    '萍乡': {'lat': 27.61672, 'lon': 113.85353},
    '营口': {'lat': 40.66472, 'lon': 122.23176},
    '葫芦岛': {'lat': 40.75243, 'lon': 120.83552},
    '蚌埠': {'lat': 32.94083, 'lon': 117.36083},
    '衡水': {'lat': 37.73908, 'lon': 115.68348},
    '衡阳': {'lat': 26.88946, 'lon': 112.61888},
    '衢州': {'lat': 28.95944, 'lon': 118.86861},
    '襄阳': {'lat': 32.0422, 'lon': 112.14479},
    '西宁': {'lat': 36.62554, 'lon': 101.75739},
    '西安': {'lat': 34.33778, 'lon': 108.70261},
    '许昌': {'lat': 34.03189, 'lon': 113.86299},
    '贵港': {'lat': 23.11603, 'lon': 109.59472},
    '贵阳': {'lat': 26.58333, 'lon': 106.71667},
    '贺州': {'lat': 24.40357, 'lon': 111.56675},
    '资阳': {'lat': 30.12108, 'lon': 104.64811},
    '赣州': {'lat': 25.84664, 'lon': 114.9326},
    '赤峰': {'lat': 42.26833, 'lon': 118.96361},
    '辽源': {'lat': 42.88545, 'lon': 125.1367},
    '辽阳': {'lat': 41.27194, 'lon': 123.17306},
    '达州': {'lat': 31.2106, 'lon': 107.46308},
    '运城': {'lat': 35.02306, 'lon': 110.99278},
    '连云港': {'lat': 34.59845, 'lon': 119.21556},
    '通化': {'lat': 41.71972, 'lon': 125.92639},
    '通辽': {'lat': 43.6125, 'lon': 122.26528},
    '遂宁': {'lat': 30.50802, 'lon': 105.57332},
    '遵义': {'lat': 27.68667, 'lon': 106.90722},
    '邢台': {'lat': 37.06217, 'lon': 114.49272},
    '那曲': {'lat': 31.0476, 'lon': 83.9694},
    '邯郸': {'lat': 36.60999, 'lon': 114.48764},
    '邵阳': {'lat': 27.23818, 'lon': 111.46214},
    '郑州': {'lat': 34.75778, 'lon': 113.64861},
    '郴州': {'lat': 25.8, 'lon': 113.03333},
    '鄂尔多斯': {'lat': 39.6086, 'lon': 109.78157},
    '鄂州': {'lat': 30.39607, 'lon': 114.88655},
    '酒泉': {'lat': 39.74318, 'lon': 98.51736},
    '重庆': {'lat': 29.56026, 'lon': 106.55771},
    '金华': {'lat': 29.10678, 'lon': 119.64421},
    '金昌': {'lat': 38.50062, 'lon': 102.19379},
    '钦州': {'lat': 21.98247, 'lon': 108.65061},
    '铁岭': {'lat': 42.29306, 'lon': 123.84139},
    '铜仁': {'lat': 27.71722, 'lon': 109.18528},
    '铜川': {'lat': 34.8988, 'lon': 108.95056},
    '铜陵': {'lat': 30.95, 'lon': 117.78333},
    '银川': {'lat': 38.46806, 'lon': 106.27306},
    '锦州': {'lat': 41.10778, 'lon': 121.14167},
    '镇江': {'lat': 32.21086, 'lon': 119.45508},
    '长春': {'lat': 43.88, 'lon': 125.32278},
    '长沙': {'lat': 28.19874, 'lon': 112.97087},
    '长治': {'lat': 34.86595, 'lon': 110.39567},
    '阜新': {'lat': 42.01556, 'lon': 121.65889},
    '阜阳': {'lat': 32.9, 'lon': 115.81667},
    '防城港': {'lat': 21.76945, 'lon': 108.35661},
    '阳江': {'lat': 21.85563, 'lon': 111.96272},
    '阳泉': {'lat': 37.8575, 'lon': 113.56333},
    '阿克苏': {'lat': 41.18418, 'lon': 80.27921},
    '阿勒泰': {'lat': 47.8505, 'lon': 88.13286},
    '陇南': {'lat': 33.39791, 'lon': 104.91703},
    '随州': {'lat': 31.71111, 'lon': 113.36306},
    '雅安': {'lat': 29.98521, 'lon': 102.999},
    '青岛': {'lat': 36.06488, 'lon': 120.38042},
    '鞍山': {'lat': 41.12361, 'lon': 122.99},
    '韶关': {'lat': 24.8, 'lon': 113.58333},
    '香港': {'lat': 22.27832, 'lon': 114.17469},
    '马鞍山': {'lat': 31.6704, 'lon': 118.5074},
    '驻马店': {'lat': 32.97944, 'lon': 114.02944},
    '高雄': {'lat': 22.61626, 'lon': 120.31333},
    '鸡西': {'lat': 45.29322, 'lon': 130.96217},
    '鹤壁': {'lat': 35.73231, 'lon': 114.28616},
    '鹤岗': {'lat': 47.34727, 'lon': 130.29033},
    '鹰潭': {'lat': 28.23333, 'lon': 117.0},
    '黄冈': {'lat': 30.45143, 'lon': 114.87035},
    '黄山': {'lat': 29.71139, 'lon': 118.3125},
    '黄石': {'lat': 30.24706, 'lon': 115.04814},
    '黑河': {'lat': 50.24413, 'lon': 127.49016},
    '齐齐哈尔': {'lat': 47.33922, 'lon': 123.96154},
    '龙岩': {'lat': 25.07485, 'lon': 117.01775},
}

def weathercode_to_cn(code):
    """将 WMO 天气代码转为中文描述。"""
    try:
        code = int(code)
    except (TypeError, ValueError):
        return '多云'
    return WMO_WEATHER.get(code, '多云')


def detect_city(text):
    """从文本（工程名称等）中识别内置城市名，返回城市中文名；无则返回空串。"""
    if not text:
        return ''
    # 按名称长度降序匹配，避免短名误匹配
    for city in sorted(CITY_COORDS.keys(), key=len, reverse=True):
        if city in text:
            return city
    return ''


def fetch_weather(city, date_str=''):
    """
    通过 Open-Meteo 获取指定城市指定日期的天气与气温（无需API Key）。
    返回 {'am': 上午天气, 'pm': 下午天气, 'high': 最高, 'low': 最低, 'avg': 平均}，
    失败时返回 None（由调用方优雅降级）。
    """
    try:
        coords = CITY_COORDS.get(city)
        if not coords:
            return None
        dt = parse_date(date_str) or datetime.date.today()
        start = dt.isoformat()
        url = ('https://api.open-meteo.com/v1/forecast?latitude=%s&longitude=%s'
               '&daily=weather_code,temperature_2m_max,temperature_2m_min,weathercode'
               '&hourly=weather_code&timezone=Asia%%2FShanghai'
               '&start_date=%s&end_date=%s' % (coords['lat'], coords['lon'], start, start))
        with urllib.request.urlopen(url, timeout=12) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        daily = data.get('daily') or {}
        times = daily.get('time') or []
        if not times:
            return None
        idx = 0
        try:
            idx = times.index(start)
        except ValueError:
            pass
        # 兼容新旧字段名
        codes = daily.get('weather_code') or daily.get('weathercode') or []
        high_arr = daily.get('temperature_2m_max') or []
        low_arr = daily.get('temperature_2m_min') or []
        code = codes[idx] if idx < len(codes) else None
        high = high_arr[idx] if idx < len(high_arr) else None
        low = low_arr[idx] if idx < len(low_arr) else None

        # 用逐小时天气拆分上午/下午
        am_code = pm_code = code
        hourly = data.get('hourly') or {}
        htimes = hourly.get('time') or []
        hcodes = hourly.get('weather_code') or []
        if htimes and hcodes:
            am_codes, pm_codes = [], []
            for t, c in zip(htimes, hcodes):
                if t.startswith(start):
                    hh = int(t[11:13])
                    if 8 <= hh < 14:
                        am_codes.append(c)
                    elif 14 <= hh < 20:
                        pm_codes.append(c)
            if am_codes:
                am_code = max(set(am_codes), key=am_codes.count)
            if pm_codes:
                pm_code = max(set(pm_codes), key=pm_codes.count)

        high_v = '' if high is None else str(int(round(float(high))))
        low_v = '' if low is None else str(int(round(float(low))))
        avg = ''
        if high is not None and low is not None:
            avg = str(int(round((float(high) + float(low)) / 2)))
        return {'am': weathercode_to_cn(am_code), 'pm': weathercode_to_cn(pm_code),
                'high': high_v, 'low': low_v, 'avg': avg}
    except Exception:
        return None


def get_template_path():
    """获取模板文件路径。"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, '..', 'assets', '监理日志.docx')
    return os.path.normpath(template_path)


def get_chars_per_line():
    """计算内容单元格每行可放多少个中文字符。"""
    available_width_pt = (CELL_WIDTH_TWIPS - CELL_MARGIN_TWIPS) / 20.0
    return int(available_width_pt / CHAR_WIDTH_PT)


def estimate_physical_lines(paragraphs_text):
    """
    估算编号段落列表在单元格中占据的物理行数（考虑自动换行）。
    
    Args:
        paragraphs_text: 段落文本列表（不含编号前缀，如["3号楼浇筑...", "钢筋绑扎..."]）
    
    Returns:
        总物理行数
    """
    chars_per_line = get_chars_per_line()
    first_line_chars = chars_per_line - FIRST_LINE_INDENT_CHARS - NUMBER_PREFIX_CHARS
    follow_line_chars = chars_per_line
    
    total_lines = 0
    for text in paragraphs_text:
        text_len = len(text)
        if text_len <= first_line_chars:
            total_lines += 1
        else:
            remaining = text_len - first_line_chars
            total_lines += 1 + (remaining + follow_line_chars - 1) // follow_line_chars
    
    return total_lines


def split_sentences(text):
    """将文本按句号、分号、问号、感叹号分割为句子列表。"""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 先按换行分割，再按句末标点分割
    sentences = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # 按中文句末标点分割（保留标点）
        parts = re.split(r'(?<=[。；！？])', line)
        for p in parts:
            p = p.strip()
            if p:
                sentences.append(p)
    return sentences


def score_sentence(sentence, category):
    """计算句子的重要性分数，包含关键词越多分数越高。"""
    score = 0
    keywords = SECTION_KEYWORDS.get(category, [])
    for kw in keywords:
        if kw in sentence:
            score += 1
    # 包含数字（工程量、时间、编号等）的句子更重要
    if re.search(r'\d+', sentence):
        score += 1
    # 较长的句子通常包含更多信息
    if len(sentence) > 20:
        score += 1
    return score


def compress_sentence(sentence):
    """对单个句子进行轻度压缩，去除冗余修饰词，保留关键信息。"""
    # 去除常见冗余表述
    redundants = [
        '的话', '的话呢', '的话吧', '的话啊',
        '其实', '实际上', '事实上', '基本上', '大致上',
        '可以说是', '可以说', '应该是', '可能是',
        '非常', '十分', '相当', '比较', '较为',
        '及时地', '有效地', '顺利地', '成功地',
        '进行了', '开展了', '实施了', '落实了',
    ]
    for r in redundants:
        sentence = sentence.replace(r, '')
    # 合并连续空格
    sentence = re.sub(r'\s+', ' ', sentence)
    return sentence.strip()


def summarize_content(text, category):
    """
    智能摘要精简：在不丢失关键信息的前提下，将内容压缩到板块字数限制内。
    
    策略：
    1. 按句子分割
    2. 计算每句重要性分数
    3. 优先保留高分句子
    4. 对保留的长句子进行轻度压缩
    5. 最后确保总字数不超限
    """
    if not text or not text.strip():
        return ''
    
    limits = SECTION_LIMITS[category]
    max_chars = limits['max_chars']
    
    # 如果原始内容已在限制内，直接返回
    if len(text) <= max_chars:
        return text
    
    sentences = split_sentences(text)
    if not sentences:
        return text[:max_chars]
    
    # 计算每句分数
    scored = [(score_sentence(s, category), i, s) for i, s in enumerate(sentences)]
    
    # 按分数降序选择句子，直到接近字数限制
    scored_sorted = sorted(scored, key=lambda x: (-x[0], x[1]))
    
    selected = []
    total_chars = 0
    for score, idx, sentence in scored_sorted:
        if total_chars + len(sentence) <= max_chars:
            selected.append((idx, sentence))
            total_chars += len(sentence)
        elif total_chars < max_chars * 0.7 and score > 0:
            # 还有较多空间时，即使单句超限也尝试压缩后加入
            compressed = compress_sentence(sentence)
            if total_chars + len(compressed) <= max_chars:
                selected.append((idx, compressed))
                total_chars += len(compressed)
    
    # 按原始顺序排列
    selected.sort(key=lambda x: x[0])
    result = ''.join(s for _, s in selected)
    
    # 如果仍超限，对结果逐句压缩
    if len(result) > max_chars:
        result_sentences = split_sentences(result)
        compressed_all = [compress_sentence(s) for s in result_sentences]
        result = ''.join(compressed_all)
    
    # 最终截断保护（保留句末标点）
    if len(result) > max_chars:
        result = result[:max_chars-1] + '。'
    
    return result


def optimize_and_number(text, category):
    """
    对输入文本进行优化表达、智能摘要、按段落编号，并控制物理行数不超过最大行数。
    
    流程：原始文本 → 优化表达 → 智能摘要 → 去除旧编号 → 物理行数控制（迭代压缩） → 重新编号
    """
    if not text or not text.strip():
        return ''
    
    limits = SECTION_LIMITS[category]
    max_lines = limits['max_lines']
    
    # 先做表达优化
    text = optimize_expression(text, category)
    
    # 智能摘要精简（应用字数限制）
    text = summarize_content(text, category)
    
    # 按句子分割
    raw_items = split_sentences(text)
    if not raw_items:
        raw_items = [text.strip()]
    
    # 去除已有的编号前缀
    cleaned_items = []
    for item in raw_items:
        item = re.sub(r'^\s*[\(（]?\d+[\)）、\.]\s*', '', item)
        item = re.sub(r'^\s*[一二三四五六七八九十]+、\s*', '', item)
        item = item.strip()
        if item:
            if not item.endswith(('。', '；', '！', '？')):
                item += '。'
            cleaned_items.append(item)
    
    if not cleaned_items:
        return text.strip()
    
    # 物理行数控制：迭代压缩直到不超过max_lines
    # 策略1：合并相邻短句子（阈值逐步放宽，避免过度合并）
    for merge_threshold in [25, 40, 55, 75, 100]:
        if estimate_physical_lines(cleaned_items) <= max_lines:
            break
        merged = []
        i = 0
        while i < len(cleaned_items):
            current = cleaned_items[i]
            # 仅当当前句子较短且与下一句合并后不超过100字时才合并
            if (len(current) < merge_threshold and i + 1 < len(cleaned_items)
                    and len(current) + len(cleaned_items[i+1]) <= 100):
                current = current + cleaned_items[i+1]
                i += 2
            else:
                i += 1
            merged.append(current)
        cleaned_items = merged
    
    # 策略2：对长句子进行轻度压缩（去除冗余修饰词）
    if estimate_physical_lines(cleaned_items) > max_lines:
        cleaned_items = [compress_sentence(s) for s in cleaned_items]
    
    # 策略3：再次合并短句子（压缩后可能变短）
    if estimate_physical_lines(cleaned_items) > max_lines:
        merged = []
        i = 0
        while i < len(cleaned_items):
            current = cleaned_items[i]
            if (len(current) < 50 and i + 1 < len(cleaned_items)
                    and len(current) + len(cleaned_items[i+1]) <= 120):
                current = current + cleaned_items[i+1]
                i += 2
            else:
                i += 1
            merged.append(current)
        cleaned_items = merged
    
    # 策略4：如果仍超限，删除重要性分数最低的句子（最多删除1/4，且至少保留2句）
    if estimate_physical_lines(cleaned_items) > max_lines and len(cleaned_items) > 2:
        max_delete = max(1, len(cleaned_items) // 4)
        deleted = 0
        while (estimate_physical_lines(cleaned_items) > max_lines 
               and deleted < max_delete 
               and len(cleaned_items) > 2):
            # 找分数最低的句子删除
            scored = [(score_sentence(s, category), i) for i, s in enumerate(cleaned_items)]
            scored_sorted = sorted(scored, key=lambda x: (x[0], x[1]))
            del_idx = scored_sorted[0][1]
            cleaned_items.pop(del_idx)
            deleted += 1
    
    # 最终截断保护：如果仍超限，截断最后一个句子
    if estimate_physical_lines(cleaned_items) > max_lines and cleaned_items:
        chars_per_line = get_chars_per_line()
        # 计算允许的最大总字符数（粗略估算）
        max_total_chars = max_lines * chars_per_line - FIRST_LINE_INDENT_CHARS * len(cleaned_items) - NUMBER_PREFIX_CHARS * len(cleaned_items)
        current_total = sum(len(s) for s in cleaned_items)
        if current_total > max_total_chars:
            # 从最后一个句子开始截断
            excess = current_total - max_total_chars
            last = cleaned_items[-1]
            if len(last) > excess + 1:
                cleaned_items[-1] = last[:len(last)-excess-1] + '。'
            else:
                cleaned_items.pop()
    
    # 重新编号
    optimized = [f'{i}、{item}' for i, item in enumerate(cleaned_items, 1)]
    return '\n'.join(optimized)


def optimize_expression(text, category):
    """根据内容类别优化表达，使语言更专业规范。"""
    if category == 'construction':
        text = text.replace('做了', '完成')
        text = text.replace('搞了', '实施')
        text = text.replace('弄了', '完成')
        text = text.replace('开始', '已开始')
        text = text.replace('在做', '正在进行')
        text = text.replace('浇混凝土', '浇筑混凝土')
        text = text.replace('扎钢筋', '绑扎钢筋')
        text = text.replace('支模', '安装模板')
        text = text.replace('拆模', '拆除模板')
        text = text.replace('挖土', '土方开挖')
        text = text.replace('回填', '土方回填')
    elif category == 'supervision':
        text = text.replace('看了', '巡视检查')
        text = text.replace('查了', '检查')
        text = text.replace('验收了', '组织验收')
        text = text.replace('旁站了', '旁站监理')
        text = text.replace('说了', '口头指示')
        text = text.replace('要求整改', '要求施工单位整改')
        text = text.replace('要求停工', '要求施工单位停工')
        text = text.replace('要求返工', '要求施工单位返工')
        text = text.replace('开会', '召开会议')
        text = text.replace('签字', '签署')
    elif category == 'problem':
        text = text.replace('有问题', '存在问题')
        text = text.replace('不合格', '不符合要求')
        text = text.replace('让改', '责令整改')
        text = text.replace('要求改', '要求整改')
        text = text.replace('已改', '已整改完成')
        text = text.replace('整改了', '已整改完成')
        text = text.replace('罚款', '按合同约定处理')
    return text


def set_paragraph_format(paragraph, line_spacing_pt=16, first_line_indent_chars=2):
    """
    设置段落格式：固定行距和首行缩进。
    
    Args:
        paragraph: 段落XML元素 (w:p)
        line_spacing_pt: 行距（磅），默认16磅
        first_line_indent_chars: 首行缩进字符数，默认2字符
    """
    # 获取或创建段落属性
    pPr = paragraph.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph.insert(0, pPr)
    
    # 设置行距：固定值 line_spacing_pt 磅
    # line 值 = 磅数 × 20（单位为二十分之一磅）
    # lineRule="exact" 表示固定值
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    
    # 设置首行缩进：first_line_indent_chars 字符
    # firstLineChars 值 = 字符数 × 100（单位为百分之一字符）
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(int(first_line_indent_chars * 100)))
    # 同时设置 firstLine（以二十分之一磅为单位），兼容旧版阅读器
    # 假设一个中文字符约22磅（小四号字），这里用一个合理的估算
    ind.set(qn('w:firstLine'), str(int(first_line_indent_chars * 22 * 20)))


def set_cell_text_formatted(cell, text, line_spacing_pt=16, first_line_indent_chars=2):
    """
    设置单元格文本，应用专业排版格式（固定行距、首行缩进）。
    
    Args:
        cell: 单元格XML元素
        text: 文本内容（多行用\n分隔）
        line_spacing_pt: 行距（磅），默认16磅
        first_line_indent_chars: 首行缩进字符数，默认2字符
    """
    # 清除现有段落内容，但保留第一个段落
    paragraphs = cell.findall(qn('w:p'))
    if not paragraphs:
        return
    
    first_p = paragraphs[0]
    for p in paragraphs[1:]:
        p.getparent().remove(p)
    
    # 获取第一个run的格式作为模板
    template_rPr = None
    runs = first_p.findall(qn('w:r'))
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            template_rPr = deepcopy(rPr)
    
    # 清除第一个段落的所有run
    for r in first_p.findall(qn('w:r')):
        first_p.remove(r)
    
    # 按行添加文本，每行一个段落
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            new_p = OxmlElement('w:p')
            cell.append(new_p)
            target_p = new_p
        else:
            target_p = first_p
        
        # 设置段落格式（固定行距、首行缩进）
        set_paragraph_format(target_p, line_spacing_pt, first_line_indent_chars)
        
        # 创建run并添加文本
        new_r = OxmlElement('w:r')
        if template_rPr is not None:
            new_r.append(deepcopy(template_rPr))
        t = OxmlElement('w:t')
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        new_r.append(t)
        target_p.append(new_r)


def set_cell_text(cell, text):
    """设置单元格文本（简单模式，不应用排版格式，用于表头、日期等短字段）。"""
    paragraphs = cell.findall(qn('w:p'))
    if not paragraphs:
        return
    first_p = paragraphs[0]
    for p in paragraphs[1:]:
        p.getparent().remove(p)
    template_rPr = None
    runs = first_p.findall(qn('w:r'))
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            template_rPr = deepcopy(rPr)
    for r in first_p.findall(qn('w:r')):
        first_p.remove(r)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            new_p = deepcopy(first_p)
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            cell.append(new_p)
            target_p = new_p
        else:
            target_p = first_p
        new_r = target_p.makeelement(qn('w:r'), {})
        if template_rPr is not None:
            new_r.append(deepcopy(template_rPr))
        t = new_r.makeelement(qn('w:t'), {})
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        new_r.append(t)
        target_p.append(new_r)


def make_page_break_paragraph():
    """创建一个分页符段落（w:br type=page），用于让每份日志另起一页。"""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def save_document(doc, output_path):
    """保存文档，自动创建输出目录。"""
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    doc.save(output_path)


def fill_diary_table(table, project_name='', date='', weather_am='', weather_pm='',
                     temp_high='', temp_low='', temp_avg='', city='',
                     construction='', supervision='', problem='', other='',
                     diary_no_year='', diary_no_seq='', supervisor='', chief=''):
    """
    在指定表格中填充一份监理日志内容（含优化表达、智能摘要、编号与专业排版）。
    单日期与多日期模式共用。
    """
    # 当日监理工作情况未提供时，使用通用表述（不允许留空）
    supervision = (supervision or '').strip() or DEFAULT_SUPERVISION

    # 天气/气温：未提供任何天气气温字段且能识别城市时，按该城市自动获取
    if not (weather_am or weather_pm or temp_high or temp_low or temp_avg):
        city_name = city or detect_city(project_name)
        if city_name:
            w = fetch_weather(city_name, date)
            if w:
                weather_am = w['am']
                weather_pm = w['pm']
                temp_high = w['high']
                temp_low = w['low']
                temp_avg = w['avg']

    # 优化表达、智能摘要、编号核心内容
    construction_opt = optimize_and_number(construction, 'construction')
    supervision_opt = optimize_and_number(supervision, 'supervision')
    problem_opt = optimize_and_number(problem, 'problem')
    # 其他有关事项也做摘要精简，但不编号
    other_summary = summarize_content(other.strip(), 'other') if other else ''

    trs = table._tbl.tr_lst

    # 行1: 工程名称、编号
    if len(trs) > 1:
        tcs = trs[1].tc_lst
        if project_name and len(tcs) >= 2:
            set_cell_text(tcs[1], project_name)
        if diary_no_year and len(tcs) >= 4:
            set_cell_text(tcs[3], diary_no_year)
        if diary_no_seq and len(tcs) >= 6:
            set_cell_text(tcs[5], diary_no_seq)

    # 行2-3: 日期、星期、天气、气温
    if len(trs) > 3:
        tcs2 = trs[2].tc_lst
        tcs3 = trs[3].tc_lst
        if date and len(tcs2) >= 2:
            set_cell_text(tcs2[1], date)
        weekday = get_weekday(date)
        if weekday and len(tcs3) >= 2:
            set_cell_text(tcs3[1], weekday)
        if weather_am and len(tcs2) >= 5:
            set_cell_text(tcs2[4], weather_am)
        if weather_pm and len(tcs3) >= 5:
            set_cell_text(tcs3[4], weather_pm)
        if temp_high and len(tcs2) >= 7:
            set_cell_text(tcs2[6], f'最高{temp_high}℃')
        if temp_low and len(tcs3) >= 7:
            set_cell_text(tcs3[6], f'最低{temp_low}℃')
        if temp_avg and len(tcs3) >= 8:
            set_cell_text(tcs3[7], f'{temp_avg}℃')

    # 四大业务板块：应用专业排版（固定行距16磅、首行缩进2字符）
    # 行5: 当日施工进展情况
    if len(trs) > 5 and construction_opt:
        set_cell_text_formatted(trs[5].tc_lst[0], construction_opt,
                                 line_spacing_pt=16, first_line_indent_chars=2)
    # 行7: 当日监理工作情况
    if len(trs) > 7 and supervision_opt:
        set_cell_text_formatted(trs[7].tc_lst[0], supervision_opt,
                                 line_spacing_pt=16, first_line_indent_chars=2)
    # 行9: 当日存在的问题及处理情况（未提供时填写"无"）
    if len(trs) > 9:
        problem_fill = problem_opt if problem_opt else '无'
        set_cell_text_formatted(trs[9].tc_lst[0], problem_fill,
                                 line_spacing_pt=16, first_line_indent_chars=2)
    # 行11: 其他有关事项（未提供时填写"无"）
    if len(trs) > 11:
        other_fill = other_summary if other_summary else '无'
        set_cell_text_formatted(trs[11].tc_lst[0], other_fill,
                                 line_spacing_pt=16, first_line_indent_chars=2)

    # 行12: 签字
    if len(trs) > 12:
        tcs12 = trs[12].tc_lst
        if supervisor and len(tcs12) >= 2:
            set_cell_text(tcs12[1], supervisor)
        if chief and len(tcs12) >= 4:
            set_cell_text(tcs12[3], chief)


def generate_diary(output_path, project_name='', date='', weather_am='', weather_pm='',
                   temp_high='', temp_low='', temp_avg='', city='',
                   construction='', supervision='', problem='', other='',
                   diary_no_year='', diary_no_seq='', supervisor='', chief=''):
    """
    生成监理日志文件（单日期，基于15列新版模板，含专业排版控制）。
    """
    template_path = get_template_path()
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'模板文件不存在: {template_path}')

    doc = Document(template_path)
    fill_diary_table(doc.tables[0], project_name=project_name, date=date,
                     weather_am=weather_am, weather_pm=weather_pm,
                     temp_high=temp_high, temp_low=temp_low, temp_avg=temp_avg,
                     city=city,
                     construction=construction, supervision=supervision,
                     problem=problem, other=other,
                     diary_no_year=diary_no_year, diary_no_seq=diary_no_seq,
                     supervisor=supervisor, chief=chief)
    save_document(doc, output_path)
    return output_path


def generate_diary_multi(output_path, entries):
    """
    生成多日期监理日志文件：将多天的日志写入同一份docx，
    每份日志之间插入分页符，各自另起一页。

    Args:
        output_path: 输出docx路径
        entries: 日志参数列表，每个元素为一天日志的关键字参数字典，
                 字段与 generate_diary 的入参一致。
    """
    if not entries:
        raise ValueError('entries 不能为空')

    template_path = get_template_path()
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'模板文件不存在: {template_path}')

    doc = Document(template_path)
    # 保存一份未填充的模板表格，供后续日期复用
    pristine_tbl = deepcopy(doc.tables[0]._tbl)

    # 填充第一天日志（使用模板自带表格）
    fill_diary_table(doc.tables[0], **entries[0])

    # 定位模板末尾的"注："段落，作为后续插入位置（让"注："与节属性保持在最后）
    body = doc.element.body
    note_para = None
    for child in body:
        if child.tag == qn('w:p'):
            texts = child.findall('.//' + qn('w:t'))
            txt = ''.join(t.text or '' for t in texts)
            if '注：' in txt:
                note_para = child
                break

    # 其余日期：复制模板表格，并在每份日志前插入分页符
    from docx.table import Table
    for entry in entries[1:]:
        new_tbl_el = deepcopy(pristine_tbl)
        if note_para is not None:
            note_para.addprevious(new_tbl_el)
            new_tbl_el.addprevious(make_page_break_paragraph())
        else:
            body.append(make_page_break_paragraph())
            body.append(new_tbl_el)
        fill_diary_table(Table(new_tbl_el, doc), **entry)

    save_document(doc, output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='监理日志生成器（含专业排版控制）')
    parser.add_argument('-o', '--output', required=True, help='输出文件路径')
    parser.add_argument('--project', default='', help='工程名称')
    parser.add_argument('--date', default='', help='日期')
    parser.add_argument('--weather-am', default='', help='上午天气')
    parser.add_argument('--weather-pm', default='', help='下午天气')
    parser.add_argument('--temp-high', default='', help='最高气温')
    parser.add_argument('--temp-low', default='', help='最低气温')
    parser.add_argument('--temp-avg', default='', help='平均温度')
    parser.add_argument('--construction', default='', help='当日施工进展情况（≤600字自动摘要）')
    parser.add_argument('--supervision', default='', help='当日监理工作情况（≤260字自动摘要）')
    parser.add_argument('--problem', default='', help='当日存在的问题及处理情况（≤160字自动摘要）')
    parser.add_argument('--other', default='', help='其他有关事项（≤120字自动摘要）')
    parser.add_argument('--diary-no-year', default='', help='监理日志编号-年份')
    parser.add_argument('--diary-no-seq', default='', help='监理日志编号-序号')
    parser.add_argument('--supervisor', default='', help='监理人员签字')
    parser.add_argument('--chief', default='', help='总监或总代签阅')
    parser.add_argument('--construction-file', default='', help='从文件读取当日施工进展情况')
    parser.add_argument('--supervision-file', default='', help='从文件读取当日监理工作情况')
    parser.add_argument('--problem-file', default='', help='从文件读取当日存在的问题及处理情况')
    parser.add_argument('--entries-file', default='', help='多日期JSON配置文件路径（JSON数组，每项为一天日志的参数）')
    parser.add_argument('--entries', default='', help='多日期JSON字符串（JSON数组，每项为一天日志的参数）')
    parser.add_argument('--city', default='', help='城市名称（如"厦门"；未提供时自动从工程名称识别，用于自动获取天气与气温）')

    args = parser.parse_args()

    # 多日期模式：从JSON文件或JSON字符串读取多天日志，写为同一份分页docx
    if args.entries_file or args.entries:
        import json
        raw = ''
        if args.entries_file:
            with open(args.entries_file, 'r', encoding='utf-8-sig') as f:
                raw = f.read()
        else:
            raw = args.entries
        data = json.loads(raw)
        if not isinstance(data, list) or not data:
            raise ValueError('--entries/--entries-file 必须是非空JSON数组，每个元素为一天的日志参数')

        entry_keys = ['project', 'date', 'weather_am', 'weather_pm', 'temp_high',
                      'temp_low', 'temp_avg', 'city', 'construction', 'supervision',
                      'problem', 'other', 'diary_no_year', 'diary_no_seq',
                      'supervisor', 'chief']
        file_keys = {'construction_file': 'construction',
                     'supervision_file': 'supervision',
                     'problem_file': 'problem'}
        entries = []
        for item in data:
            if not isinstance(item, dict):
                raise ValueError('entries 的每个元素必须是JSON对象')
            entry = {}
            for k in entry_keys:
                # JSON 中的 project 对应函数入参 project_name
                key = 'project_name' if k == 'project' else k
                entry[key] = str(item.get(k, ''))
            # 兼容同时提供 project_name 的情况
            if not entry.get('project_name'):
                entry['project_name'] = str(item.get('project_name', ''))
            # 支持从文件读取核心内容（长文本场景）
            for fk, target in file_keys.items():
                fp = item.get(fk, '')
                if fp and os.path.exists(fp):
                    with open(fp, 'r', encoding='utf-8-sig') as f:
                        entry[target] = f.read()
            entries.append(entry)

        output = generate_diary_multi(args.output, entries)
        print(f'✅ 监理日志已生成（共{len(entries)}个日期，每个日期另起一页）: {output}')
        return

    construction = args.construction
    supervision = args.supervision
    problem = args.problem
    
    if args.construction_file and os.path.exists(args.construction_file):
        with open(args.construction_file, 'r', encoding='utf-8-sig') as f:
            construction = f.read()
    if args.supervision_file and os.path.exists(args.supervision_file):
        with open(args.supervision_file, 'r', encoding='utf-8-sig') as f:
            supervision = f.read()
    if args.problem_file and os.path.exists(args.problem_file):
        with open(args.problem_file, 'r', encoding='utf-8-sig') as f:
            problem = f.read()
    
    output = generate_diary(
        output_path=args.output,
        project_name=args.project,
        date=args.date,
        weather_am=args.weather_am,
        weather_pm=args.weather_pm,
        temp_high=args.temp_high,
        temp_low=args.temp_low,
        temp_avg=args.temp_avg,
        city=args.city,
        construction=construction,
        supervision=supervision,
        problem=problem,
        other=args.other,
        diary_no_year=args.diary_no_year,
        diary_no_seq=args.diary_no_seq,
        supervisor=args.supervisor,
        chief=args.chief
    )
    
    print(f'✅ 监理日志已生成: {output}')


if __name__ == '__main__':
    main()

