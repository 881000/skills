#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
监理通知单生成器
基于标准监理通知单模板生成格式一致的docx文件，
并对事由（主题）进行简短化优化、对内容（正文）进行专业表达优化和段落编号。
"""
import argparse
import os
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn


def get_template_path():
    """获取模板文件路径"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, '..', 'assets', '监理通知单.docx')
    return os.path.normpath(template_path)


def optimize_subject(subject):
    """
    优化事由（主题），使其简短、规范。
    统一格式为"关于XXX的事宜"，去除冗余表述。
    """
    if not subject or not subject.strip():
        return ''

    text = subject.strip()

    # 去除已有的"关于""的事宜""的通知""问题"等套话前缀后缀，提取核心
    text = re.sub(r'^关于\s*', '', text)
    text = re.sub(r'\s*(的事宜|的相关事宜|事宜|的通知|通知|的问题|问题)$', '', text)
    # 去除末尾残留的"整改通知""安全通知"等组合词中的"通知"
    text = re.sub(r'(整改|安全|质量|进度)\s*通知$', r'\1', text)
    text = text.strip('，。；：、 ')

    # 限制长度，超过30字则截取核心部分
    if len(text) > 30:
        text = text[:30]

    # 根据结尾是否含"的"决定拼接方式，避免"的XX的事宜"拗口
    if text.endswith('的'):
        return f'关于{text}事宜'
    return f'关于{text}的事宜'


def optimize_and_number_content(text):
    """
    对内容（正文）进行专业表达优化并按段落编号。
    输出格式：1、…。2、…。以此类推，每条独立成段。
    """
    if not text or not text.strip():
        return ''

    # 按换行、中文分号、句号切分为独立条目
    raw_items = []
    lines = [l.strip() for l in text.replace('\r\n', '\n').split('\n') if l.strip()]

    for line in lines:
        # 对每行再按中文分号、句号分割（保留分隔符用于判断）
        parts = re.split(r'(?<=[；。])', line)
        for p in parts:
            p = p.strip()
            if p:
                raw_items.append(p)

    if not raw_items:
        raw_items = [text.strip()]

    # 去除已有的编号前缀
    cleaned_items = []
    for item in raw_items:
        # 去除 "1、" "1." "（1）" "(1)" "1)" 等编号前缀
        item = re.sub(r'^\s*[\(（]?\d+[\)）、\.]\s*', '', item)
        # 去除 "一、" "二、" 等中文数字编号
        item = re.sub(r'^\s*[一二三四五六七八九十]+、\s*', '', item)
        # 去除 "首先""其次""然后""最后"等序列词（含后面的顿号/逗号）
        item = re.sub(r'^\s*(首先|其次|然后|最后|接着|此外|另外)\s*[，,、]?\s*', '', item)
        item = item.strip()
        if item:
            # 去除末尾的冒号、逗号等非句末标点，再补句号
            item = item.rstrip('，,：:、 ')
            # 确保以句号结尾
            if not item.endswith(('。', '；', '！', '？')):
                item += '。'
            cleaned_items.append(item)

    if not cleaned_items:
        return text.strip()

    # 专业表达优化
    optimized = []
    for i, item in enumerate(cleaned_items, 1):
        item = optimize_content_expression(item)
        optimized.append(f'{i}、{item}')

    return '\n'.join(optimized)


def optimize_content_expression(text):
    """
    对监理通知单正文内容进行专业表达优化，
    使语言符合工程监理通知的正式、规范语气。
    """
    # 口语化→工程术语
    replacements = [
        ('做了', '已完成'),
        ('搞了', '已实施'),
        ('弄了', '已完成'),
        ('在做', '正在进行'),
        ('开始做', '已开始'),
        ('浇混凝土', '浇筑混凝土'),
        ('扎钢筋', '绑扎钢筋'),
        ('支模', '安装模板'),
        ('拆模', '拆除模板'),
        ('挖土', '土方开挖'),
        ('回填土', '土方回填'),
        ('砌墙', '砌筑墙体'),
        ('抹灰', '墙面抹灰'),
        ('刷漆', '涂刷涂料'),
        ('防水', '防水层施工'),
        ('保温', '保温层施工'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)

    # 监理动作规范化
    supervision_replacements = [
        ('看了', '巡视检查发现'),
        ('查了', '检查发现'),
        ('发现有', '发现存在'),
        ('有问题', '存在问题'),
        ('不合格', '不符合设计及规范要求'),
        ('没做', '未按要求实施'),
        ('没有做', '未按要求实施'),
        ('没按', '未按'),
        ('让整改', '责令限期整改'),
        ('要求整改', '责令限期整改'),
        ('要求改', '责令限期整改'),
        ('马上改', '立即整改'),
        ('赶紧改', '立即整改'),
        ('停工', '暂停施工'),
        ('返工', '返工处理'),
        ('罚款', '按合同约定处理'),
        ('说了', '口头指示'),
        ('开会', '召开专题会议'),
    ]
    for old, new in supervision_replacements:
        text = text.replace(old, new)

    # 通知语气强化：在问题描述类条目中适当增加要求性表述
    # （仅当条目未包含"要求""责令""应""须"等指令词时，不自动添加，避免过度修改）

    return text


def set_cell_text(cell, text, keep_format=True):
    """
    设置单元格文本，尽量保留原有格式。
    支持多行文本（\n分隔），每行独立成段。
    """
    paragraphs = cell.findall(qn('w:p'))
    if not paragraphs:
        return

    # 保留第一个段落，删除其余
    first_p = paragraphs[0]
    for p in paragraphs[1:]:
        p.getparent().remove(p)

    # 获取第一个run的格式作为模板
    template_rPr = None
    runs = first_p.findall(qn('w:r'))
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            template_rPr = deepcopy(rPr)

    # 清除第一个段落的所有run
    for r in first_p.findall(qn('w:r')):
        first_p.remove(r)

    # 按行添加文本
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            new_p = deepcopy(first_p)
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            cell.append(new_p)
            target_p = new_p
        else:
            target_p = first_p

        new_r = target_p.makeelement(qn('w:r'), {})
        if template_rPr is not None:
            new_r.append(deepcopy(template_rPr))

        t = new_r.makeelement(qn('w:t'), {})
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        new_r.append(t)
        set_run_font_songti(new_r)
        target_p.append(new_r)


def set_run_font_songti(r):
    """
    设置run的字体为宋体（中文、英文、高位ansi、复杂脚本均设为宋体）。
    """
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    for rFonts in rPr.findall(qn('w:rFonts')):
        rPr.remove(rFonts)
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): '宋体',
        qn('w:eastAsia'): '宋体',
        qn('w:hAnsi'): '宋体',
        qn('w:cs'): '宋体',
    })
    rPr.insert(0, rFonts)


def set_content_paragraph_format(p):
    """
    设置正文内容段落格式：
    - 段前0行，段后0行，1.5倍行距
    - 左缩进0.5字符，右缩进0.5字符，首行缩进2字符
    注意：spacing 和 ind 必须按 OOXML 规范插入到 pPr 的正确位置
    （snapToGrid 之后、jc 之前），否则 Word/WPS 解析会异常。
    """
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = p.makeelement(qn('w:pPr'), {})
        p.insert(0, pPr)

    # 移除已有的 spacing 和 ind
    for tag in ('w:spacing', 'w:ind'):
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)

    # 创建 spacing：段前0、段后0、行距1.5倍
    spacing = pPr.makeelement(qn('w:spacing'), {
        qn('w:before'): '0',
        qn('w:after'): '0',
        qn('w:line'): '360',
        qn('w:lineRule'): 'auto',
    })

    # 创建 ind：左缩进0.5字符，右缩进0.5字符，首行缩进2字符
    ind = pPr.makeelement(qn('w:ind'), {
        qn('w:leftChars'): '50',
        qn('w:left'): '120',
        qn('w:rightChars'): '50',
        qn('w:right'): '120',
        qn('w:firstLineChars'): '200',
        qn('w:firstLine'): '480',
    })

    # 找到插入锚点：pPr 中第一个顺序在 ind 之后的元素
    # OOXML pPr 子元素顺序：... snapToGrid(21), spacing(22), ind(23),
    # contextualSpacing(24), mirrorIndents(25), suppressOverlap(26),
    # jc(27), textDirection(28), textAlignment(29) ...
    anchor = None
    for child in pPr:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('contextualSpacing', 'mirrorIndents', 'suppressOverlap',
                   'jc', 'textDirection', 'textAlignment', 'textboxTightWrap',
                   'outlineLvl', 'divId', 'cnfStyle', 'rPr'):
            anchor = child
            break

    if anchor is not None:
        anchor.addprevious(spacing)
        spacing.addnext(ind)
    else:
        pPr.append(spacing)
        pPr.append(ind)


def fill_notice_body_cell(cell, contractor='', subject='', content='', deadline=''):
    """
    填写通知单正文单元格（行2，包含"致""事由""内容"三部分）。
    保留原有标签段落结构，精确替换各部分内容。
    """
    paragraphs = cell.findall(qn('w:p'))
    if len(paragraphs) < 3:
        return

    # --- 段落0: 致：___（施工项目经理部）---
    p0 = paragraphs[0]
    # 清除现有run，重建
    for r in p0.findall(qn('w:r')):
        p0.remove(r)
    # 获取原格式模板（从已删除的run中无法获取，使用单元格默认）
    # 直接创建新run
    r0 = p0.makeelement(qn('w:r'), {})
    t0 = r0.makeelement(qn('w:t'), {})
    if contractor:
        t0.text = f'致：{contractor}（施工项目经理部）'
    else:
        t0.text = '致：                 （施工项目经理部）'
    t0.set(qn('xml:space'), 'preserve')
    r0.append(t0)
    set_run_font_songti(r0)
    p0.append(r0)

    # --- 段落1: 事由：关于XXX的事宜 ---
    p1 = paragraphs[1]
    for r in p1.findall(qn('w:r')):
        p1.remove(r)
    r1 = p1.makeelement(qn('w:r'), {})
    t1 = r1.makeelement(qn('w:t'), {})
    if subject:
        t1.text = f'事由：{subject}'
    else:
        t1.text = '事由：关于...的相关事宜'
    t1.set(qn('xml:space'), 'preserve')
    r1.append(t1)
    set_run_font_songti(r1)
    p1.append(r1)

    # --- 段落2: 内容： ---
    p2 = paragraphs[2]
    for r in p2.findall(qn('w:r')):
        p2.remove(r)
    r2 = p2.makeelement(qn('w:r'), {})
    t2 = r2.makeelement(qn('w:t'), {})
    t2.text = '内容：'
    t2.set(qn('xml:space'), 'preserve')
    r2.append(t2)
    set_run_font_songti(r2)
    p2.append(r2)

    # --- 段落3及以后：清除，准备填入编号内容 ---
    for p in paragraphs[3:]:
        p.getparent().remove(p)

    # 在"内容："段落后插入编号内容段落
    if content:
        content_lines = content.split('\n')
        for line in content_lines:
            new_p = deepcopy(p2)
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # 设置正文段落格式：段前0行、段后0行、1.5倍行距、左缩进0.5字符、首行缩进2字符
            set_content_paragraph_format(new_p)
            r_new = new_p.makeelement(qn('w:r'), {})
            t_new = r_new.makeelement(qn('w:t'), {})
            t_new.text = line
            t_new.set(qn('xml:space'), 'preserve')
            r_new.append(t_new)
            set_run_font_songti(r_new)
            new_p.append(r_new)
            # 插入到p2之后
            p2.addnext(new_p)
            p2 = new_p  # 下一段插在当前段之后

    # 在正文内容之后追加整改要求段落
    if content:
        if deadline:
            rectify_text = (
                f'针对上述问题，请贵单位严格按照相关规范及标准要求进行整改落实。'
                f'整改工作须于{deadline}前完成，并将书面整改回复报送项目监理部。'
                f'经监理部复查验收合格后，方可恢复施工。'
            )
        else:
            rectify_text = (
                '针对上述问题，请贵单位严格按照相关规范及标准要求进行整改落实。'
                '并将书面整改回复报送项目监理部。'
                '经监理部复查验收合格后，方可恢复施工。'
            )
        rectify_p = deepcopy(p2)
        for r in rectify_p.findall(qn('w:r')):
            rectify_p.remove(r)
        set_content_paragraph_format(rectify_p)
        r_rectify = rectify_p.makeelement(qn('w:r'), {})
        t_rectify = r_rectify.makeelement(qn('w:t'), {})
        t_rectify.text = rectify_text
        t_rectify.set(qn('xml:space'), 'preserve')
        r_rectify.append(t_rectify)
        set_run_font_songti(r_rectify)
        rectify_p.append(r_rectify)
        p2.addnext(rectify_p)


def fill_notice_no_cell(cell, year='', seq=''):
    """
    填写通知单编号单元格。
    原格式："监理[      ]通知     号"
    填入后："监理[ year ]通知 seq 号"
    """
    paragraphs = cell.findall(qn('w:p'))
    if not paragraphs:
        return
    p = paragraphs[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)

    r = p.makeelement(qn('w:r'), {})
    t = r.makeelement(qn('w:t'), {})
    year_str = year if year else '      '
    seq_str = seq if seq else '     '
    t.text = f'监理[{year_str}]通知{seq_str}号'
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    set_run_font_songti(r)
    p.append(r)


def fill_signature_cell(cell, engineer='', date_str=''):
    """
    填写签字区单元格。
    包含：项目监理机构（盖章）、总/专业监理工程师（签字）、日期
    """
    paragraphs = cell.findall(qn('w:p'))
    if len(paragraphs) < 3:
        return

    # 段落0: 项目监理机构（盖章）： - 保持不变
    # 段落1: 总/专业监理工程师（签字）：
    p1 = paragraphs[1]
    if engineer:
        for r in p1.findall(qn('w:r')):
            p1.remove(r)
        r = p1.makeelement(qn('w:r'), {})
        t = r.makeelement(qn('w:t'), {})
        t.text = f'总/专业监理工程师（签字）：{engineer}'
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        set_run_font_songti(r)
        p1.append(r)

    # 段落2: 日 期：      年    月    日
    p2 = paragraphs[2]
    if date_str:
        for r in p2.findall(qn('w:r')):
            p2.remove(r)
        r = p2.makeelement(qn('w:r'), {})
        t = r.makeelement(qn('w:t'), {})
        t.text = f'日 期：{date_str}'
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        set_run_font_songti(r)
        p2.append(r)


def set_document_font_songti(doc):
    """
    遍历整个文档（段落+表格），将所有run的字体统一设置为宋体。
    包括模板中原有的标题、标签、签收区、表尾注释等。
    """
    # 文档正文段落
    for p in doc.paragraphs:
        for r in p.runs:
            set_run_font_songti(r._r)
    # 表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_font_songti(r._r)


def generate_notice(output_path, project_name='', notice_year='', notice_seq='',
                    contractor='', subject='', content='', engineer='', date_str='',
                    deadline=''):
    """
    生成监理通知单文件。

    Args:
        output_path: 输出文件路径
        project_name: 工程名称
        notice_year: 通知单编号-年份（如"2026"）
        notice_seq: 通知单编号-序号（如"001"）
        contractor: 致（施工单位/项目经理部名称）
        subject: 事由（原始主题描述，将自动优化为简短格式）
        content: 内容（原始正文描述，将自动优化并编号）
        engineer: 总/专业监理工程师（签字）
        date_str: 日期（如"2026年8月21日"）
    """
    template_path = get_template_path()
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'模板文件不存在: {template_path}')

    doc = Document(template_path)
    table = doc.tables[0]
    trs = table._tbl.tr_lst

    # 优化事由和内容
    subject_opt = optimize_subject(subject)
    content_opt = optimize_and_number_content(content)

    # 行1: 工程名称 + 编号
    if len(trs) > 1:
        tcs = trs[1].tc_lst
        # tc[0]="工程名称："标签, tc[1]=空(span2,填值), tc[2]=编号
        # 注意：由于tc[1]是gridSpan=2，实际tc列表为[tc0, tc1, tc2]
        if project_name and len(tcs) >= 2:
            set_cell_text(tcs[1], project_name)
        # 编号单元格：最后一个tc
        if len(tcs) >= 3:
            fill_notice_no_cell(tcs[-1], notice_year, notice_seq)

    # 行2: 致/事由/内容
    if len(trs) > 2:
        tcs = trs[2].tc_lst
        if tcs:
            fill_notice_body_cell(tcs[0], contractor, subject_opt, content_opt, deadline)

    # 行3: 签字区
    if len(trs) > 3:
        tcs = trs[3].tc_lst
        if tcs:
            fill_signature_cell(tcs[0], engineer, date_str)

    # 全局统一字体为宋体（中文、英文）
    set_document_font_songti(doc)

    # 保存
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='监理通知单生成器')
    parser.add_argument('-o', '--output', required=True, help='输出docx文件路径')
    parser.add_argument('--project', default='', help='工程名称')
    parser.add_argument('--notice-year', default='', help='通知单编号-年份（如2026）')
    parser.add_argument('--notice-seq', default='', help='通知单编号-序号（如001）')
    parser.add_argument('--contractor', default='', help='致（施工单位/项目经理部名称）')
    parser.add_argument('--subject', default='', help='事由（原始主题描述，将自动优化为简短格式）')
    parser.add_argument('--content', default='', help='内容（原始正文，将自动优化并按1、2、编号）')
    parser.add_argument('--engineer', default='', help='总/专业监理工程师（签字）')
    parser.add_argument('--date', default='', help='日期（如2026年8月21日）')
    parser.add_argument('--deadline', default='', help='整改截止日期（如2026年9月5日，将自动追加整改要求段落）')

    # 支持从文件读取长内容
    parser.add_argument('--content-file', default='', help='从文件读取正文内容')

    args = parser.parse_args()

    content = args.content
    if args.content_file and os.path.exists(args.content_file):
        with open(args.content_file, 'r', encoding='utf-8') as f:
            content = f.read()

    output = generate_notice(
        output_path=args.output,
        project_name=args.project,
        notice_year=args.notice_year,
        notice_seq=args.notice_seq,
        contractor=args.contractor,
        subject=args.subject,
        content=content,
        engineer=args.engineer,
        date_str=args.date,
        deadline=args.deadline
    )

    print(f'✅ 监理通知单已生成: {output}')


if __name__ == '__main__':
    main()
